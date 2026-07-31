from __future__ import annotations

import re
import tempfile
import zipfile
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.request import Request, urlopen

from lxml import etree


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}
PLACEHOLDER = re.compile(r"X{2,}", re.I)
UNRESOLVED_MARKER = re.compile(r"20XX|X{2,}", re.I)
PART_RE = re.compile(r"word/(document|header\d+|footer\d+|footnotes|endnotes)\.xml")
COMMENTS_PART = "word/comments.xml"


def _paragraph_text(paragraph) -> str:
    parts: list[str] = []
    for node in paragraph.iter():
        if node.tag in {f"{{{W}}}t", f"{{{W}}}delText"} and node.text:
            parts.append(node.text)
        elif node.tag == f"{{{W}}}tab":
            parts.append("\t")
        elif node.tag in {f"{{{W}}}br", f"{{{W}}}cr"}:
            parts.append("\n")
    return "".join(parts)


def _highlight_text(paragraph) -> str:
    values: list[str] = []
    for run in paragraph.xpath(".//w:r", namespaces=NS):
        highlight = run.find("w:rPr/w:highlight", namespaces=NS)
        if highlight is not None and highlight.get(f"{{{W}}}val") not in (None, "none", "auto"):
            values.append("".join(run.xpath(".//w:t/text()", namespaces=NS)))
    return "".join(values).strip()


def _paragraph_text_without_highlight(paragraph) -> str:
    parts: list[str] = []
    for run in paragraph.xpath(".//w:r", namespaces=NS):
        highlight = run.find("w:rPr/w:highlight", namespaces=NS)
        if highlight is not None and highlight.get(f"{{{W}}}val") not in (None, "none", "auto"):
            run_text = "".join(run.xpath(".//w:t/text()|.//w:delText/text()", namespaces=NS))
            parts.extend(match.group(0) for match in PLACEHOLDER.finditer(run_text))
            continue
        for node in run.iter():
            if node.tag in {f"{{{W}}}t", f"{{{W}}}delText"} and node.text:
                parts.append(node.text)
            elif node.tag == f"{{{W}}}tab":
                parts.append("\t")
            elif node.tag in {f"{{{W}}}br", f"{{{W}}}cr"}:
                parts.append("\n")
    text = "".join(parts)
    return re.sub(r"（\s*）|\(\s*\)", "", text)


def _comment_texts(archive: zipfile.ZipFile) -> dict[str, str]:
    """Read Word comments without requiring python-docx comment support.

    Word stores comments in a separate OOXML part and anchors them in the
    document with ``commentRangeStart``/``commentReference`` elements.  The
    communication template now uses those comments as the source-of-truth
    annotations, so inventory records retain the exact comment text and IDs.
    Older yellow-highlight templates simply return an empty mapping.
    """
    if COMMENTS_PART not in archive.namelist():
        return {}
    root = etree.fromstring(archive.read(COMMENTS_PART))
    return {
        str(comment.get(f"{{{W}}}id")): "".join(
            comment.xpath(".//w:t/text()", namespaces=NS)
        ).strip()
        for comment in root.xpath(".//w:comment", namespaces=NS)
        if comment.get(f"{{{W}}}id") is not None
    }


def extract_word_comments(path: Path) -> dict[str, str]:
    """Return all Word comments keyed by OOXML comment ID."""
    with zipfile.ZipFile(path) as archive:
        return _comment_texts(archive)


def _paragraph_comment_ids(paragraph) -> list[str]:
    ids: list[str] = []
    for node in paragraph.xpath(
        ".//w:commentRangeStart|.//w:commentReference",
        namespaces=NS,
    ):
        value = node.get(f"{{{W}}}id")
        if value is not None and value not in ids:
            ids.append(str(value))
    return ids


def inventory_template(path: Path) -> list[dict]:
    records: list[dict] = []
    with zipfile.ZipFile(path) as zf:
        comments = _comment_texts(zf)
        for part in sorted(name for name in zf.namelist() if PART_RE.fullmatch(name)):
            root = etree.fromstring(zf.read(part))
            short = Path(part).stem.upper()
            for p_index, paragraph in enumerate(root.xpath(".//w:p", namespaces=NS), 1):
                raw = _paragraph_text(paragraph)
                context = re.sub(r"\s+", " ", raw).strip()
                markers = list(PLACEHOLDER.finditer(context))
                highlight = _highlight_text(paragraph)
                in_table = bool(paragraph.xpath("ancestor::w:tbl", namespaces=NS))
                comment_ids = _paragraph_comment_ids(paragraph)
                comment_texts = [comments[item] for item in comment_ids if item in comments]
                for occurrence, match in enumerate(markers, 1):
                    records.append({
                        "location_id": f"{short}-P{p_index:04d}-X{occurrence:02d}",
                        "record_type": "占位符",
                        "part": part,
                        "paragraph_index": p_index,
                        "occurrence_index": occurrence,
                        "marker": match.group(0),
                        "context": context,
                        "in_table": in_table,
                        "comment_ids": comment_ids,
                        "comment_texts": comment_texts,
                    })
                if highlight and not markers:
                    records.append({
                        "location_id": f"{short}-P{p_index:04d}-H01",
                        "record_type": "黄色标注内容块",
                        "part": part,
                        "paragraph_index": p_index,
                        "occurrence_index": 1,
                        "marker": "黄色标注",
                        "context": context,
                        "in_table": in_table,
                        "comment_ids": comment_ids,
                        "comment_texts": comment_texts,
                    })
                # Newer communication templates use comments on headings and
                # table lead-ins that contain no literal XXX marker.  Keep
                # those anchors in the inventory so their source instruction
                # is still available for mapping/audit (the final clean
                # template simply ignores the synthetic C location during
                # replacement).
                if comment_ids and not markers and not highlight:
                    records.append({
                        "location_id": f"{short}-P{p_index:04d}-C01",
                        "record_type": "批注内容块",
                        "paragraph_index": p_index,
                        "occurrence_index": 1,
                        "marker": "",
                        "context": context,
                        "in_table": in_table,
                        "comment_ids": comment_ids,
                        "comment_texts": comment_texts,
                    })
    return records


def document_paragraph_texts(path: Path) -> list[tuple[int, str]]:
    with zipfile.ZipFile(path) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    return [
        (index, _paragraph_text(paragraph))
        for index, paragraph in enumerate(root.xpath(".//w:p", namespaces=NS), 1)
    ]


def _is_highlighted(run) -> bool:
    highlight = run.find("w:rPr/w:highlight", namespaces=NS)
    return highlight is not None and highlight.get(f"{{{W}}}val") not in (None, "none", "auto")


def _run_text_nodes(run):
    return run.xpath(".//w:t|.//w:delText", namespaces=NS)


def _run_text(run) -> str:
    return "".join(node.text or "" for node in _run_text_nodes(run))


def _clear_run_text(run) -> None:
    for node in _run_text_nodes(run):
        node.text = ""


def _remove_highlight(run) -> None:
    for highlight in run.xpath("./w:rPr/w:highlight", namespaces=NS):
        highlight.getparent().remove(highlight)


def _set_yellow_highlight(run) -> None:
    properties = run.find("w:rPr", namespaces=NS)
    if properties is None:
        properties = etree.Element(f"{{{W}}}rPr")
        run.insert(0, properties)
    for highlight in properties.findall("w:highlight", namespaces=NS):
        properties.remove(highlight)
    highlight = etree.SubElement(properties, f"{{{W}}}highlight")
    highlight.set(f"{{{W}}}val", "yellow")


def _remove_paragraph_highlights(paragraph) -> None:
    for highlight in paragraph.xpath(".//w:highlight", namespaces=NS):
        parent = highlight.getparent()
        if parent is not None:
            parent.remove(highlight)


def _set_run_text(run, value: str) -> None:
    nodes = _run_text_nodes(run)
    if not nodes:
        run.append(etree.Element(f"{{{W}}}t"))
        nodes = _run_text_nodes(run)
    nodes[0].text = value
    for node in nodes[1:]:
        node.text = ""


def _normal_runs(paragraph):
    return [run for run in paragraph.xpath("./w:r", namespaces=NS) if not _is_highlighted(run)]


def _set_paragraph_text(
    paragraph,
    value: str,
    remove_highlight: bool = False,
    *,
    style_run=None,
) -> None:
    runs = paragraph.xpath("./w:r", namespaces=NS)
    if not runs:
        style_run = etree.SubElement(paragraph, f"{{{W}}}r")
        etree.SubElement(style_run, f"{{{W}}}t")
        runs = [style_run]
    if style_run is None:
        normal = _normal_runs(paragraph)
        style_run = normal[0] if normal else runs[0]
    _set_run_text(style_run, value)
    for run in runs:
        if run is not style_run:
            _clear_run_text(run)
    if remove_highlight:
        _remove_paragraph_highlights(paragraph)


def _visible_chunks(paragraph):
    return [(_run_text(run), _is_highlighted(run), run) for run in paragraph.xpath("./w:r", namespaces=NS)]


def _strip_yellow_annotation(paragraph) -> str:
    """Remove yellow instructions and only their adjacent annotation parentheses."""
    chunks = _visible_chunks(paragraph)
    visible: list[str] = []
    for index, (text, highlighted, _run) in enumerate(chunks):
        if highlighted:
            continue
        if index > 0 and chunks[index - 1][1]:
            text = re.sub(r"^\s*[）)]", "", text)
        if index + 1 < len(chunks) and chunks[index + 1][1]:
            text = re.sub(r"[（(]\s*$", "", text)
        visible.append(text)
    return re.sub(r"（\s*）|\(\s*\)", "", "".join(visible))


def _replace_yellow_annotation(paragraph, value: str) -> None:
    chunks = _visible_chunks(paragraph)
    first_highlight = next((i for i, (_, highlighted, _) in enumerate(chunks) if highlighted), None)
    if first_highlight is None:
        _set_paragraph_text(paragraph, value, True)
        return
    prefix_parts: list[str] = []
    suffix_parts: list[str] = []
    for index, (text, highlighted, _run) in enumerate(chunks):
        if highlighted:
            continue
        if index < first_highlight:
            if index + 1 < len(chunks) and chunks[index + 1][1]:
                text = re.sub(r"[（(]\s*$", "", text)
            prefix_parts.append(text)
        else:
            if index > 0 and chunks[index - 1][1]:
                text = re.sub(r"^\s*[）)]", "", text)
            suffix_parts.append(text)
    prefix = "".join(prefix_parts)
    suffix = "".join(suffix_parts)
    if value and prefix and not prefix.endswith(("：", ":", "；", ";", "。")):
        prefix += "："
    target = next((run for _text, highlighted, run in chunks if highlighted), None)
    _set_paragraph_text(paragraph, prefix + str(value) + suffix, True, style_run=target)


def _replace_method_heading(paragraph, value: str) -> None:
    runs = paragraph.xpath("./w:r", namespaces=NS)
    normal = [run for run in runs if not _is_highlighted(run)]
    if len(normal) >= 2:
        _set_run_text(normal[1], str(value))
        for run in runs:
            if _is_highlighted(run) or run not in normal[:2]:
                _clear_run_text(run)
                _remove_highlight(run)
        _remove_paragraph_highlights(paragraph)
    else:
        _set_paragraph_text(paragraph, str(value), True)


def _replace_placeholders_preserving_runs(paragraph, items, replacements) -> None:
    ordered = sorted(items, key=lambda item: item["occurrence_index"])
    values = iter(str(replacements[item["location_id"]]) for item in ordered)
    runs = paragraph.xpath("./w:r", namespaces=NS)
    for index, run in enumerate(runs):
        original = _run_text(run)
        if not original:
            continue
        matches = list(PLACEHOLDER.finditer(original))
        if not matches:
            continue
        parts: list[str] = []
        cursor = 0
        for match in matches:
            parts.extend((original[cursor:match.start()], next(values, match.group(0))))
            cursor = match.end()
        parts.append(original[cursor:])
        if index > 0 and _is_highlighted(runs[index - 1]):
            parts[0] = re.sub(r"^\s*[）)]", "", parts[0])
        if index + 1 < len(runs) and _is_highlighted(runs[index + 1]):
            parts[-1] = re.sub(r"[（(]\s*$", "", parts[-1])
        _set_run_text(run, "".join(parts))
    for index, run in enumerate(runs):
        if _is_highlighted(run):
            continue
        text = _run_text(run)
        if index > 0 and _is_highlighted(runs[index - 1]):
            text = re.sub(r"^\s*[）)]", "", text)
        if index + 1 < len(runs) and _is_highlighted(runs[index + 1]):
            text = re.sub(r"[（(]\s*$", "", text)
        _set_run_text(run, text)
    for run in runs:
        if _is_highlighted(run):
            _clear_run_text(run)
    _remove_paragraph_highlights(paragraph)


def _set_cell_text(cell, value: str) -> None:
    paragraphs = cell.xpath("./w:p", namespaces=NS)
    if not paragraphs:
        paragraphs = [etree.SubElement(cell, f"{{{W}}}p")]
    _set_paragraph_text(paragraphs[0], value, True)
    for paragraph in paragraphs[1:]:
        cell.remove(paragraph)


def _split_paragraph_and_highlight_markers(paragraph) -> bool:
    runs = paragraph.xpath(".//w:r", namespaces=NS)
    run_texts = [_run_text(run) for run in runs]
    combined = "".join(run_texts)
    matches = list(UNRESOLVED_MARKER.finditer(combined))
    if not matches:
        return False
    offset = 0
    for run, text in zip(runs, run_texts, strict=True):
        start = offset
        end = offset + len(text)
        offset = end
        if not text:
            continue
        boundaries = {0, len(text)}
        for match in matches:
            overlap_start = max(start, match.start())
            overlap_end = min(end, match.end())
            if overlap_start < overlap_end:
                boundaries.add(overlap_start - start)
                boundaries.add(overlap_end - start)
        ordered = sorted(boundaries)
        chunks = [
            (text[left:right], start + left, start + right)
            for left, right in zip(ordered, ordered[1:])
            if left < right
        ]
        parent = run.getparent()
        if parent is None:
            continue
        position = parent.index(run)
        for value, global_start, global_end in chunks:
            clone = deepcopy(run)
            _set_run_text(clone, value)
            highlighted = any(
                match.start() <= global_start and global_end <= match.end()
                for match in matches
            )
            if highlighted:
                _set_yellow_highlight(clone)
            else:
                _remove_highlight(clone)
            parent.insert(position, clone)
            position += 1
        parent.remove(run)
    return True


def _table_coordinates(root, paragraph) -> tuple[int, int, int] | None:
    cells = paragraph.xpath("ancestor::w:tc[1]", namespaces=NS)
    tables = paragraph.xpath("ancestor::w:tbl[1]", namespaces=NS)
    if not cells or not tables:
        return None
    cell = cells[0]
    table = tables[0]
    all_tables = root.xpath(".//w:tbl", namespaces=NS)
    table_index = next(
        (index for index, candidate in enumerate(all_tables, 1) if candidate is table),
        None,
    )
    if table_index is None:
        return None
    rows = table.xpath("./w:tr", namespaces=NS)
    for row_index, row in enumerate(rows, 1):
        row_cells = row.xpath("./w:tc", namespaces=NS)
        for column_index, candidate in enumerate(row_cells, 1):
            if candidate is cell:
                return table_index, row_index, column_index
    return None


def highlight_unresolved_placeholders(path: Path) -> list[dict[str, Any]]:
    """Highlight unresolved markers and return exact Word-part locations."""
    with zipfile.ZipFile(path) as archive:
        items = archive.infolist()
        contents = {info.filename: archive.read(info.filename) for info in items}

    findings: list[dict[str, Any]] = []
    changed = False
    for part in sorted(name for name in contents if PART_RE.fullmatch(name)):
        root = etree.fromstring(contents[part])
        short = Path(part).stem.upper()
        paragraphs = root.xpath(".//w:p", namespaces=NS)
        part_changed = False
        for paragraph_index, paragraph in enumerate(paragraphs, 1):
            context = re.sub(r"\s+", " ", _paragraph_text(paragraph)).strip()
            matches = list(UNRESOLVED_MARKER.finditer(context))
            if not matches:
                continue
            coordinates = _table_coordinates(root, paragraph)
            for occurrence, match in enumerate(matches, 1):
                if coordinates is None:
                    location_id = (
                        f"{short}-P{paragraph_index:04d}-X{occurrence:02d}"
                    )
                    location_type = "段落"
                    table_index = row_index = column_index = ""
                else:
                    table_index, row_index, column_index = coordinates
                    location_id = (
                        f"{short}-T{table_index:02d}-R{row_index:02d}"
                        f"-C{column_index:02d}-X{occurrence:02d}"
                    )
                    location_type = "表格单元格"
                findings.append(
                    {
                        "location_id": location_id,
                        "location_type": location_type,
                        "part": part,
                        "paragraph_index": paragraph_index,
                        "occurrence_index": occurrence,
                        "context": context,
                        "current_text": match.group(0),
                        "table_index": table_index,
                        "row_index": row_index,
                        "column_index": column_index,
                    }
                )
            part_changed = (
                _split_paragraph_and_highlight_markers(paragraph) or part_changed
            )
        if part_changed:
            contents[part] = etree.tostring(
                root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone=True,
            )
            changed = True

    if changed:
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as output:
            for info in items:
                output.writestr(info, contents[info.filename])
    return findings


def replace_image_markers(path: Path) -> None:
    """Replace QCC image markers with the actual trademark image."""
    from docx import Document
    from docx.shared import Inches

    document = Document(str(path))
    changed = False
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                marker = next((p.text.strip() for p in cell.paragraphs if "__QCC_IMAGE__" in p.text), "")
                if not marker:
                    continue
                url = marker.split("__QCC_IMAGE__", 1)[1].strip()
                try:
                    request = Request(url, headers={"User-Agent": "asset-appraisal/1.0"})
                    with urlopen(request, timeout=15) as response:
                        data = response.read()
                    paragraph = cell.paragraphs[0]
                    paragraph.text = ""
                    paragraph.add_run().add_picture(BytesIO(data), width=Inches(0.42))
                except Exception:
                    cell.text = "图样"
                changed = True
    if changed:
        document.save(str(path))


def unresolved_placeholders(path: Path) -> list[str]:
    """Return template placeholders still present after a fill operation."""
    markers: set[str] = set()
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if not PART_RE.fullmatch(name):
                continue
            text = archive.read(name).decode("utf-8", errors="ignore")
            markers.update(re.findall(r"20XX|X{2,}", text))
    return sorted(markers)


def replace_report_number_year(path: Path, year: Any) -> None:
    """Synchronize every report-number year, including legacy literal years.

    Some template occurrences use ``20XX`` while older pages contain a literal
    ``2024``.  They are the same report-number field and must not diverge.
    """
    year_text = str(year or "").strip()
    if not re.fullmatch(r"(?:19|20)\d{2}", year_text):
        return
    with zipfile.ZipFile(path) as zin:
        items = zin.infolist()
        contents = {info.filename: zin.read(info.filename) for info in items}
    changed = False
    for name, data in list(contents.items()):
        if not PART_RE.fullmatch(name):
            continue
        root = etree.fromstring(data)
        for paragraph in root.xpath(".//w:p", namespaces=NS):
            text = _paragraph_text(paragraph)
            if "银信评报字" not in text:
                continue
            updated = re.sub(
                r"(银信评报字[（(])(?:20XX|(?:19|20)\d{2})([）)]第\s*\d+号)",
                rf"\g<1>{year_text}\g<2>",
                text,
            )
            if updated != text:
                runs = paragraph.xpath("./w:r", namespaces=NS)
                if runs:
                    _set_paragraph_text(paragraph, updated, False, style_run=runs[0])
                    changed = True
        if changed:
            contents[name] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    if not changed:
        return
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in items:
            zout.writestr(info, contents[info.filename])


def _fill_tables(root, table_replacements: dict[int, list[list[str]]]) -> None:
    tables = root.xpath(".//w:tbl", namespaces=NS)
    for table_index, matrix in table_replacements.items():
        if table_index >= len(tables):
            raise ValueError(f"Word 表格编号不存在：{table_index}")
        table = tables[table_index]
        rows = table.xpath("./w:tr", namespaces=NS)
        if len(matrix) > len(rows):
            if not rows:
                raise ValueError(f"Word 表格 {table_index} 没有可复制的行")
            template_row = rows[-1]
            for _ in range(len(matrix) - len(rows)):
                table.append(deepcopy(template_row))
            rows = table.xpath("./w:tr", namespaces=NS)
        for row in rows[len(matrix):]:
            table.remove(row)
        for row_index, values in enumerate(matrix):
            cells = rows[row_index].xpath("./w:tc", namespaces=NS)
            if len(values) != len(cells):
                raise ValueError(f"Word 表格 {table_index} 第 {row_index + 1} 行列数不匹配")
            for cell, value in zip(cells, values, strict=True):
                _set_cell_text(cell, str(value))


def _set_table_column_ratios(
    root,
    table_column_ratios: dict[int, list[float]],
) -> None:
    tables = root.xpath(".//w:tbl", namespaces=NS)
    for table_index, ratios in table_column_ratios.items():
        if table_index >= len(tables):
            raise ValueError(f"Word 表格编号不存在：{table_index}")
        table = tables[table_index]
        grid_columns = table.xpath("./w:tblGrid/w:gridCol", namespaces=NS)
        if not grid_columns or len(grid_columns) != len(ratios):
            raise ValueError(f"Word 表格 {table_index} 的列宽配置与实际列数不匹配")
        total_ratio = sum(ratios)
        if total_ratio <= 0 or any(value <= 0 for value in ratios):
            raise ValueError(f"Word 表格 {table_index} 的列宽比例必须为正数")
        current_widths = [
            int(column.get(f"{{{W}}}w") or 0)
            for column in grid_columns
        ]
        total_width = sum(current_widths) or 9000
        widths = [
            max(1, round(total_width * ratio / total_ratio))
            for ratio in ratios
        ]
        widths[-1] += total_width - sum(widths)
        for column, width in zip(grid_columns, widths, strict=True):
            column.set(f"{{{W}}}w", str(width))
        for row in table.xpath("./w:tr", namespaces=NS):
            cells = row.xpath("./w:tc", namespaces=NS)
            if len(cells) != len(widths):
                continue
            for cell, width in zip(cells, widths, strict=True):
                properties = cell.find("w:tcPr", namespaces=NS)
                if properties is None:
                    properties = etree.Element(f"{{{W}}}tcPr")
                    cell.insert(0, properties)
                cell_width = properties.find("w:tcW", namespaces=NS)
                if cell_width is None:
                    cell_width = etree.SubElement(properties, f"{{{W}}}tcW")
                cell_width.set(f"{{{W}}}w", str(width))
                cell_width.set(f"{{{W}}}type", "dxa")


def fill_template(
    template: Path,
    output: Path,
    replacements: dict[str, str],
    *,
    table_replacements: dict[int, list[list[str]]] | None = None,
    table_column_ratios: dict[int, list[float]] | None = None,
    paragraph_replacements: dict[tuple[str, int], str] | None = None,
    replacement_modes: dict[str, str] | None = None,
) -> Path:
    if template.resolve() == output.resolve():
        raise ValueError("输出 Word 不能覆盖模板")
    locations = inventory_template(template)
    by_part_para: dict[tuple[str, int], list[dict]] = {}
    for item in locations:
        if item["location_id"] in replacements:
            by_part_para.setdefault((item["part"], item["paragraph_index"]), []).append(item)
    output.parent.mkdir(parents=True, exist_ok=True)
    table_replacements = table_replacements or {}
    table_column_ratios = table_column_ratios or {}
    paragraph_replacements = paragraph_replacements or {}
    replacement_modes = replacement_modes or {}
    with zipfile.ZipFile(template) as zin, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            relevant = [(key, value) for key, value in by_part_para.items() if key[0] == info.filename]
            static_replacements = [
                (key, value) for key, value in paragraph_replacements.items() if key[0] == info.filename
            ]
            has_tables = info.filename == "word/document.xml" and bool(
                table_replacements or table_column_ratios
            )
            if relevant or has_tables or static_replacements:
                root = etree.fromstring(data)
                paragraphs = root.xpath(".//w:p", namespaces=NS)
                for (_, p_index), items in relevant:
                    paragraph = paragraphs[p_index - 1]
                    if items[0]["record_type"] == "黄色标注内容块":
                        value = str(replacements[items[0]["location_id"]])
                        mode = replacement_modes.get(items[0]["location_id"], "replace_paragraph")
                        if mode == "strip_yellow_annotation":
                            if value:
                                _replace_yellow_annotation(paragraph, value)
                            else:
                                _set_paragraph_text(paragraph, _strip_yellow_annotation(paragraph), True)
                        elif mode == "strip_yellow_only":
                            if UNRESOLVED_MARKER.fullmatch(value):
                                _replace_yellow_annotation(paragraph, value)
                            else:
                                _set_paragraph_text(
                                    paragraph,
                                    _strip_yellow_annotation(paragraph),
                                    True,
                                )
                        elif mode == "replace_yellow_annotation":
                            _replace_yellow_annotation(paragraph, value)
                        elif mode == "replace_method_heading":
                            _replace_method_heading(paragraph, value)
                        elif value:
                            _set_paragraph_text(paragraph, value, True)
                        else:
                            _set_paragraph_text(paragraph, _strip_yellow_annotation(paragraph), True)
                    else:
                        _replace_placeholders_preserving_runs(paragraph, items, replacements)
                for (_, p_index), value in static_replacements:
                    _set_paragraph_text(paragraphs[p_index - 1], str(value), True)
                if has_tables:
                    _fill_tables(root, table_replacements)
                    _set_table_column_ratios(root, table_column_ratios)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(info, data)
    return output
