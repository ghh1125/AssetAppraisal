from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook


def _pt(value: Any) -> float | None:
    return round(float(value.pt), 2) if value is not None and hasattr(value, "pt") else None


def _paragraph_snapshot(paragraph, index: int) -> dict[str, Any]:
    fmt = paragraph.paragraph_format
    return {
        "index": index,
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else "",
        "alignment": str(paragraph.alignment) if paragraph.alignment is not None else "",
        "first_line_indent_pt": _pt(fmt.first_line_indent),
        "left_indent_pt": _pt(fmt.left_indent),
        "right_indent_pt": _pt(fmt.right_indent),
        "space_before_pt": _pt(fmt.space_before),
        "space_after_pt": _pt(fmt.space_after),
        "line_spacing": fmt.line_spacing,
        "runs": [
            {
                "text": run.text,
                "font": run.font.name,
                "size_pt": _pt(run.font.size),
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
            }
            for run in paragraph.runs
        ],
    }


def _table_snapshot(table, index: int) -> dict[str, Any]:
    rows = []
    for row in table.rows[:100]:
        rows.append([cell.text for cell in row.cells])
    return {
        "index": index,
        "rows": len(table.rows),
        "columns": len(table.columns),
        "style": table.style.name if table.style else "",
        "cells": rows,
    }


def word_structure_snapshot(path: Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [_paragraph_snapshot(paragraph, index) for index, paragraph in enumerate(document.paragraphs, 1)]
    tables = [_table_snapshot(table, index) for index, table in enumerate(document.tables, 1)]
    sections = [
        {
            "page_width_in": _pt(section.page_width) / 72 if section.page_width else None,
            "page_height_in": _pt(section.page_height) / 72 if section.page_height else None,
            "top_margin_in": _pt(section.top_margin) / 72 if section.top_margin else None,
            "bottom_margin_in": _pt(section.bottom_margin) / 72 if section.bottom_margin else None,
            "left_margin_in": _pt(section.left_margin) / 72 if section.left_margin else None,
            "right_margin_in": _pt(section.right_margin) / 72 if section.right_margin else None,
        }
        for section in document.sections
    ]
    return {
        "path": path.name,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs[:1000],
        "tables": tables,
        "sections": sections,
    }


def build_format_review_evidence(template_path: Path, report_path: Path) -> dict[str, Any]:
    return {
        "template": word_structure_snapshot(template_path),
        "report": word_structure_snapshot(report_path),
    }


def _report_text(path: Path, limit: int = 60000) -> str:
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        chunks.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(chunks)[:limit]


def build_semantic_review_evidence(report_path: Path, fields: dict[str, Any]) -> dict[str, Any]:
    return {"report_text": _report_text(report_path), "fields": fields}


def build_data_review_evidence(
    report_path: Path,
    fields: dict[str, Any],
    audit_path: Path,
    evidence: dict[str, Any],
) -> dict[str, Any]:
    audit_rows: list[dict[str, Any]] = []
    workbook = load_workbook(audit_path, read_only=True, data_only=True)
    sheet = workbook["填充结果"]
    headers = [cell.value for cell in next(sheet.iter_rows())]
    for row in sheet.iter_rows(min_row=2, values_only=True):
        audit_rows.append({str(key): value for key, value in zip(headers, row)})
    workbook.close()
    return {
        "report_text": _report_text(report_path),
        "fields": fields,
        "field_audit": audit_rows,
        "evidence": evidence,
    }
