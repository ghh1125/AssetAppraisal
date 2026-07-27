from __future__ import annotations

from pathlib import Path

from docx import Document


NARRATIVE_EVIDENCE_KEYWORDS = (
    "公司概况",
    "基本情况",
    "历史沿革",
    "经营范围",
    "主营业务",
    "业务模式",
    "主要产品",
    "主要服务",
    "行业",
    "市场",
    "客户",
    "供应商",
    "采购模式",
    "生产模式",
    "销售模式",
    "盈利模式",
    "竞争优势",
    "竞争劣势",
    "SWOT",
    "可比公司",
)


def read_narrative_evidence(
    path: Path,
    *,
    source_name: str = "reference_report",
    max_blocks: int = 48,
) -> list[dict[str, str]]:
    """Extract bounded, cited narrative evidence from an uploaded DOCX.

    Relevant headings are kept together with nearby body paragraphs, while
    matching table rows preserve legal-profile and business facts.  File names
    and fixed paragraph coordinates are intentionally not used.
    """
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    paragraph_scores: dict[int, int] = {}
    for index, text in enumerate(paragraphs):
        if text and any(keyword.lower() in text.lower() for keyword in NARRATIVE_EVIDENCE_KEYWORDS):
            paragraph_scores[index] = max(paragraph_scores.get(index, 0), 100)
            for offset, score in ((-1, 70), (1, 90), (2, 80), (3, 70)):
                candidate = index + offset
                if 0 <= candidate < len(paragraphs) and paragraphs[candidate]:
                    paragraph_scores[candidate] = max(
                        paragraph_scores.get(candidate, 0),
                        score,
                    )
    # The beginning of a report commonly contains the company identity and
    # transaction overview even when headings use project-specific wording.
    for index, text in enumerate(paragraphs[:60]):
        if text:
            paragraph_scores[index] = max(paragraph_scores.get(index, 0), 10)

    table_evidence: list[dict[str, str]] = []
    table_limit = min(12, max(0, max_blocks // 4))
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            text = "｜".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if not text or not any(
                keyword.lower() in text.lower()
                for keyword in (*NARRATIVE_EVIDENCE_KEYWORDS, "统一社会信用代码", "注册资本", "法定代表人")
            ):
                continue
            table_evidence.append(
                {
                    "evidence_id": f"document:{source_name}:t{table_index}:r{row_index}",
                    "text": text[:1600],
                }
            )
            if len(table_evidence) >= table_limit:
                break
        if len(table_evidence) >= table_limit:
            break

    evidence: list[dict[str, str]] = []
    paragraph_limit = max(0, max_blocks - len(table_evidence))
    ranked_paragraphs = sorted(
        paragraph_scores,
        key=lambda index: (-paragraph_scores[index], index),
    )[:paragraph_limit]
    for index in sorted(ranked_paragraphs):
        evidence.append(
            {
                "evidence_id": f"document:{source_name}:p{index + 1}",
                "text": paragraphs[index][:1600],
            }
        )
    return [*evidence, *table_evidence]


def read_paragraph_containing(path: Path, text: str) -> str:
    """Return the first body paragraph containing the configured evidence text."""
    for paragraph in Document(path).paragraphs:
        value = paragraph.text.strip()
        if text in value:
            return value
    return ""


def read_table_cell(
    path: Path,
    *,
    table_contains: str,
    row_contains: str,
    column: int,
) -> str:
    """Return one cell from the first matching DOCX table row."""
    for table in Document(path).tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if table_contains not in table_text:
            continue
        for row in table.rows:
            if row_contains in "\t".join(cell.text for cell in row.cells):
                if column >= len(row.cells):
                    raise ValueError(f"Word 表格列号不存在：{column}")
                return row.cells[column].text.strip()
    return ""


def read_table_matrix(path: Path, table_index: int) -> list[list[str]]:
    """Read a DOCX table as plain text without changing the source document."""
    tables = Document(path).tables
    if table_index < 0 or table_index >= len(tables):
        raise ValueError(f"Word 表格编号不存在：{table_index}")
    return [[cell.text.strip() for cell in row.cells] for row in tables[table_index].rows]
