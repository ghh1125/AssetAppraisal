import json
from pathlib import Path

import pytest
from docx import Document

from demo.adapters import excel
from demo.run import run_project


PROJECT = Path("demo/projects/tongfu.yaml")
REFERENCE = Path("资产评估工作流/通富昆山评估报告-S2.docx")


def _rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def test_configured_financial_tables_match_filled_s2_reference():
    config = json.loads(PROJECT.read_text(encoding="utf-8"))
    base = PROJECT.parent
    source = (base / config["sources"]["audited_financials"]).resolve()
    expected_doc = Document(REFERENCE)
    for spec, expected_index in zip(config["financial_tables"], (4, 5), strict=True):
        actual = excel.read_configured_table(source, spec)
        assert actual == _rows(expected_doc.tables[expected_index])


def test_real_run_populates_balance_and_income_tables(tmp_path: Path):
    result = run_project(PROJECT, output_dir=tmp_path, offline=True, report_date="2025-12-01")
    generated = Document(result.report_path)
    expected = Document(REFERENCE)
    assert _rows(generated.tables[4]) == _rows(expected.tables[4])
    assert _rows(generated.tables[5]) == _rows(expected.tables[5])
    fields = json.loads((tmp_path / "normalized_fields.json").read_text(encoding="utf-8"))
    assert fields["historical_income_statement_table"]["rows"][-1] == [
        "四、净利润", "14,357,065.14", "2,600,607.69", "10,028,484.84"
    ]
    assert fields["book_net_assets"] == 4598.157666
    assert fields["income_approach_value"] == 8500
    assert fields["income_increment"] == pytest.approx(3901.842334)
    assert fields["income_increment_rate"] == pytest.approx(84.85664514836581)
    assert fields["asset_approach_value"] == 6365.043162
    assert fields["asset_increment"] == pytest.approx(1766.885496)
    assert fields["asset_increment_rate"] == pytest.approx(38.42594413551368)
    assert fields["final_appraisal_value"] == 8500
    assert fields["final_value_chinese"] == "捌仟伍佰万"
    assert fields["final_value_chinese_wan"] == "捌仟伍佰"
    assert fields["appraisal_increment"] == pytest.approx(3901.842334)
    assert fields["appraisal_increment_rate"] == pytest.approx(84.85664514836581)
    assert fields["final_valuation_method"] == "收益法"
