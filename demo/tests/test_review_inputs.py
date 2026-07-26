from pathlib import Path

from docx import Document

from demo.adapters.review_inputs import build_format_review_evidence, build_semantic_review_evidence


def test_format_review_evidence_contains_template_and_report_structure(tmp_path: Path):
    template = tmp_path / "template.docx"
    report = tmp_path / "report.docx"
    for path, text in ((template, "模板"), (report, "报告")):
        document = Document()
        paragraph = document.add_paragraph(text)
        paragraph.style = "Normal"
        document.add_table(rows=1, cols=2).cell(0, 0).text = "表头"
        document.save(path)

    evidence = build_format_review_evidence(template, report)

    assert evidence["template"]["paragraph_count"] == 1
    assert evidence["report"]["paragraphs"][0]["text"] == "报告"
    assert evidence["report"]["tables"][0]["rows"] == 1


def test_semantic_review_evidence_keeps_report_text_and_fields(tmp_path: Path):
    report = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("公司概况示例")
    document.save(report)

    evidence = build_semantic_review_evidence(report, {"target_company_name": "示例公司"})

    assert "公司概况示例" in evidence["report_text"]
    assert evidence["fields"]["target_company_name"] == "示例公司"
