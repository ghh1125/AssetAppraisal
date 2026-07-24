import hashlib
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

from demo.adapters.word import fill_template, inventory_template, replace_report_number_year, unresolved_placeholders
from demo.domain.replacement import build_replacements
from demo.domain.field_validation import normalize_report_serial


def test_replaces_two_markers_and_yellow_instruction_without_touching_template(tmp_path: Path):
    template = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("甲方XXX收购乙方XXX股权")
    p = doc.add_paragraph()
    r = p.add_run("黄色业务说明")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.save(template)
    before = hashlib.sha256(template.read_bytes()).hexdigest()
    locations = inventory_template(template)
    replacements = {
        locations[0]["location_id"]: "甲公司",
        locations[1]["location_id"]: "乙公司",
        locations[2]["location_id"]: "实际生成的业务概述",
    }
    output = tmp_path / "filled.docx"
    fill_template(template, output, replacements)
    assert hashlib.sha256(template.read_bytes()).hexdigest() == before
    assert output != template
    with zipfile.ZipFile(output) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    assert "甲公司" in xml and "乙公司" in xml and "实际生成的业务概述" in xml
    assert "XXX" not in xml and "黄色业务说明" not in xml and 'w:val="yellow"' not in xml


def test_formats_split_date_fields_for_20xx_template_markers():
    locations = [
        {"location_id": "P-X01", "field_key": "valuation_date_year", "field_name": "年", "context": "20XX年XX月XX日", "marker": "XX"},
        {"location_id": "P-X02", "field_key": "valuation_date_month", "field_name": "月", "context": "20XX年XX月XX日", "marker": "XX"},
        {"location_id": "P-X03", "field_key": "valuation_date_day", "field_name": "日", "context": "20XX年XX月XX日", "marker": "XX"},
    ]
    fields = {key: "2025年06月30日" for key in ("valuation_date_year", "valuation_date_month", "valuation_date_day")}
    assert build_replacements(locations, fields) == {"P-X01": "25", "P-X02": "06", "P-X03": "30"}


def test_normalizes_full_report_number_to_template_serial():
    assert normalize_report_serial("苏正评报字（2025）第001号") == "001"
    assert normalize_report_serial("001号") == "001"
    assert normalize_report_serial("001") == "001"


def test_replaces_legacy_literal_report_years_and_leaves_no_placeholders(tmp_path: Path):
    template = tmp_path / "report-number.docx"
    doc = Document()
    doc.add_paragraph("银信评报字（2000）第001号")
    doc.add_paragraph("银信评报字（2024）第001号")
    doc.save(template)
    output = tmp_path / "filled.docx"
    fill_template(template, output, {})
    replace_report_number_year(output, "2025")
    text = "\n".join(p.text for p in Document(output).paragraphs)
    assert text == "银信评报字（2025）第001号\n银信评报字（2025）第001号"
    assert unresolved_placeholders(output) == []


def test_formats_financial_values_by_mapping_unit():
    locations = [
        {"location_id": "P-X01", "field_key": "income_approach_value", "field_name": "收益法评估值", "context": "评估值XX万元", "marker": "XX", "unit_scope": "万元"},
        {"location_id": "P-X02", "field_key": "income_increment_rate", "field_name": "收益法增值率", "context": "增值率XX%", "marker": "XX", "unit_scope": "%"},
    ]
    fields = {"income_approach_value": 8500, "income_increment_rate": 84.85664514836581}
    assert build_replacements(locations, fields) == {"P-X01": "8,500.00", "P-X02": "84.86"}


def test_avoids_repeating_suffix_already_present_after_marker():
    locations = [
        {"location_id": "P-X01", "field_key": "method", "field_name": "方法", "context": "最终采用XX法", "marker": "XX"},
        {"location_id": "Q-X01", "field_key": "subject", "field_name": "对象", "context": "被评估单位XXX价值", "marker": "XXX"},
        {"location_id": "R-X01", "field_key": "company", "field_name": "公司", "context": "收购XXX有限责任公司股权", "marker": "XXX"},
    ]
    fields = {"method": "收益法", "subject": "股东全部权益价值", "company": "通富热处理（昆山）有限公司"}
    assert list(build_replacements(locations, fields).values()) == ["收益", "股东全部权益", "通富热处理（昆山）"]


def test_writes_matrix_into_existing_word_table_and_trims_unused_rows(tmp_path: Path):
    template = tmp_path / "table-template.docx"
    doc = Document()
    table = doc.add_table(rows=5, cols=4)
    table.cell(0, 0).text = "项目"
    doc.save(template)
    matrix = [
        ["项目", "2023年度", "2024年度", "2025年1-6月"],
        ["营业收入", "46,186,357.24", "23,707,106.86", "21,652,412.25"],
        ["净利润", "14,357,065.14", "2,600,607.69", "10,028,484.84"],
    ]
    output = tmp_path / "filled-table.docx"
    fill_template(template, output, {}, table_replacements={0: matrix})
    filled = Document(output)
    assert [[cell.text for cell in row.cells] for row in filled.tables[0].rows] == matrix


def test_removes_each_highlighted_instruction_and_its_empty_parentheses(tmp_path: Path):
    template = tmp_path / "instructions.docx"
    doc = Document()
    paragraph = doc.add_paragraph("账面值XXX（")
    note = paragraph.add_run("PDF审计报告中的所有者权益/净资产")
    note.font.highlight_color = WD_COLOR_INDEX.YELLOW
    paragraph.add_run("）万元，评估值XXX（")
    note = paragraph.add_run("表格数据-收益法表格")
    note.font.highlight_color = WD_COLOR_INDEX.YELLOW
    paragraph.add_run("）万元。")
    doc.save(template)
    locations = inventory_template(template)
    output = tmp_path / "filled.docx"
    fill_template(
        template,
        output,
        {
            locations[0]["location_id"]: "4,598.16",
            locations[1]["location_id"]: "8,500.00",
        },
    )
    text = Document(output).paragraphs[0].text
    assert text == "账面值4,598.16万元，评估值8,500.00万元。"


def test_blank_yellow_replacement_preserves_non_highlighted_heading(tmp_path: Path):
    template = tmp_path / "mixed-highlight.docx"
    doc = Document()
    paragraph = doc.add_paragraph("三、评估对象和评估范围（")
    note = paragraph.add_run("优先PDF审计报告获取、表格获取")
    note.font.highlight_color = WD_COLOR_INDEX.YELLOW
    paragraph.add_run("）")
    doc.save(template)
    location = inventory_template(template)[0]
    output = tmp_path / "filled.docx"
    fill_template(template, output, {location["location_id"]: ""})
    assert Document(output).paragraphs[0].text == "三、评估对象和评估范围"


def test_replaces_configured_static_paragraphs_without_placeholders(tmp_path: Path):
    template = tmp_path / "static.docx"
    doc = Document()
    doc.add_paragraph("截止报告出具日，前海联合基金申报的账外无形资产情况如下：")
    doc.add_paragraph("1、已申请注册的商标")
    doc.save(template)
    output = tmp_path / "filled.docx"
    fill_template(
        template,
        output,
        {},
        paragraph_replacements={
            ("word/document.xml", 1): "",
            ("word/document.xml", 2): "1、商标：现有材料未见商标申报记录。",
        },
    )
    paragraphs = Document(output).paragraphs
    assert paragraphs[0].text == ""
    assert paragraphs[1].text == "1、商标：现有材料未见商标申报记录。"


def test_preserves_heading_and_body_run_format_for_method_route(tmp_path: Path):
    template = tmp_path / "method.docx"
    doc = Document()
    p = doc.add_paragraph()
    heading = p.add_run("十一、评估方法：")
    heading.bold = True
    method = p.add_run("收益法、市场法")
    method.bold = False
    note = p.add_run("（节点勾选）")
    note.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.save(template)
    location = inventory_template(template)[0]
    output = tmp_path / "filled.docx"
    fill_template(
        template,
        output,
        {location["location_id"]: "收益法、资产基础法"},
        replacement_modes={location["location_id"]: "replace_method_heading"},
    )
    filled = Document(output)
    assert filled.paragraphs[0].text == "十一、评估方法：收益法、资产基础法"
    runs = filled.paragraphs[0].runs
    assert runs[0].text == "十一、评估方法：" and runs[0].bold is True
    assert runs[1].text == "收益法、资产基础法" and runs[1].bold is not True


def test_preserves_placeholder_run_formatting(tmp_path: Path):
    template = tmp_path / "styled.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("标题：").bold = True
    value = p.add_run("XXX")
    value.font.size = Pt(12)
    doc.save(template)
    location = inventory_template(template)[0]
    output = tmp_path / "filled.docx"
    fill_template(template, output, {location["location_id"]: "通富昆山"})
    runs = Document(output).paragraphs[0].runs
    assert runs[0].bold is True and runs[0].text == "标题："
    assert runs[1].text == "通富昆山" and runs[1].font.size == Pt(12)


def test_yellow_annotation_value_keeps_heading_and_removes_parentheses(tmp_path: Path):
    template = tmp_path / "narrative.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("3、被评估单位概述").bold = True
    p.add_run("（")
    note = p.add_run("基于评估主体生成")
    note.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run("）")
    doc.save(template)
    location = inventory_template(template)[0]
    output = tmp_path / "filled.docx"
    fill_template(
        template,
        output,
        {location["location_id"]: "公司成立于2011年。"},
        replacement_modes={location["location_id"]: "replace_yellow_annotation"},
    )
    assert Document(output).paragraphs[0].text == "3、被评估单位概述：公司成立于2011年。"
