import hashlib
import json
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

import demo.pipeline as pipeline_module
from demo.pipeline import _apply_ocr_overrides_to_table, _company_profile_table, _ocr_ownership_matrix, _validated_qcc_payload, run_pipeline


def test_company_profile_table_writes_credit_code_and_profile_values():
    rows = _company_profile_table(
        {
            "credit_code": "91320000608319749X",
            "name": "示例有限公司",
            "company_type": "有限责任公司",
            "registered_capital": "1,000万元",
            "status": "存续",
        }
    )
    assert rows[0] == [
        "统一社会信用代码：91320000608319749X",
        "企业名称：示例有限公司",
    ]
    assert rows[1][0] == "类型：有限责任公司"
    assert rows[2][0] == "注册资本：1,000万元"


def test_qcc_identity_mismatch_is_rejected_instead_of_filling_wrong_profile():
    issues = []
    assert _validated_qcc_payload(
        {"profile": {"name": "另一家有限公司"}, "fields": {"commissioning_party_profile": "错误"}},
        "目标有限公司",
        "被评估单位",
        issues,
    ) == {}
    assert "企查查身份核验失败" in issues[0]


def test_ocr_ownership_matrix_joins_split_shareholder_names():
    matrix, names = _ocr_ownership_matrix(
        {
            "table_cells": [
                {"table_id": "p39-t4", "row": 2, "column": 1, "text": "上海上大热处理有"},
                {"table_id": "p39-t4", "row": 3, "column": 1, "text": "限公司"},
                {"table_id": "p39-t4", "row": 3, "column": 2, "text": "5,000,000.00"},
                {"table_id": "p39-t4", "row": 4, "column": 1, "text": "富士和机械工业（昆"},
                {"table_id": "p39-t4", "row": 5, "column": 1, "text": "5,000,000.00"},
                {"table_id": "p39-t4", "row": 6, "column": 1, "text": "山）有限公司"},
            ]
        },
        {
            "table_id": "p39-t4",
            "target_table_index": 2,
            "rows": [
                {"name_cells": [[2, 1], [3, 1]], "capital_cell": [3, 2], "percent": "50%"},
                {"name_cells": [[4, 1], [6, 1]], "capital_cell": [5, 1], "percent": "50%"},
            ],
        },
    )
    assert names == ["上海上大热处理有限公司", "富士和机械工业（昆山）有限公司"]
    assert matrix[-1] == ["合计", "合计", "10,000,000.00", "100%"]


def test_ocr_amount_overrides_keep_word_table_in_sync():
    matrix = [["其中：固定资产账面金额：", "4,993,561.04"], ["无形资产账面金额：", "0.00"]]
    spec = {"rows": [
        {"label": "其中：固定资产账面金额：", "ocr_field_key": "fixed"},
        {"label": "无形资产账面金额：", "ocr_field_key": "intangibles"},
    ]}
    assert _apply_ocr_overrides_to_table(matrix, spec, {"fixed": "5,050,511.04", "intangibles": "96,508.64"}) == [
        ["其中：固定资产账面金额：", "5,050,511.04"],
        ["无形资产账面金额：", "96,508.64"],
    ]


class FixtureOcrAdapter:
    def extract(self, pdf_path):
        assert pdf_path.suffix == ".pdf"
        return [
            {
                "page_number": 1,
                "page_count": 1,
                "blocks": [
                    {
                        "block_id": "p1-b1",
                        "block_type": "text",
                        "text": "通富昆山审计材料",
                        "confidence": 0.99,
                        "bbox": [1, 2, 3, 4],
                    }
                ],
                "tables": [],
            }
        ], []


class FixtureLlmAdapter:
    prompt_version = "yellow_narratives.test"

    def generate(self, evidence):
        assert evidence["evidence"][0]["evidence_id"] == "pdf:p1:b1"
        return {
            "company_profile_section": "基于 OCR 证据生成的公司概述。",
            "main_products": "热处理服务。",
            "ownership_history": "不应采用的 LLM 越权内容",
        }, ["模拟 LLM 返回了越权字段"]


class NoPdfEvidenceLlmAdapter:
    prompt_version = "yellow_narratives.test"

    def generate(self, evidence):
        by_id = {item["evidence_id"]: item["text"] for item in evidence["evidence"]}
        assert "field:target_company_name" in by_id
        assert "示例有限公司" in by_id["field:target_company_name"]
        return {"company_profile_section": "根据已上传结构化材料生成。"}, []


class ReferenceDocumentEvidenceLlmAdapter:
    prompt_version = "yellow_narratives.test"

    def generate(self, evidence):
        document_blocks = [
            item
            for item in evidence["evidence"]
            if item["evidence_id"].startswith("document:reference_report:")
        ]
        assert any("主营产品为工业滤波器" in item["text"] for item in document_blocks)
        return {"main_products": "主营产品为工业滤波器。"}, []


class FixtureQichachaAdapter:
    def fetch(self, company_name):
        assert "通富" in company_name
        return {
            "ownership_history": "企查查返回的历史股权沿革。",
            "industry_overview": "不应采用的 API 越权内容",
        }, []


class CapitalQichachaAdapter:
    def fetch(self, company_name):
        return {
            "profile": {
                "name": company_name,
                "registered_capital": "1,250万元",
            },
            "fields": {},
        }, []


class FixtureTemplatePageReader:
    def extract(self, template_path):
        assert template_path.suffix == ".docx"
        return ["XXX有限责任公司拟收购", "纳入评估范围的全部资产和负债。"], []


def fixture_ocr_fields(normalized, config):
    assert normalized["text_blocks"]
    return {
        "tax_rates": "增值税税率13%，企业所得税税率15%。",
        "valuation_scope": "纳入评估范围的全部资产和负债。",
    }, []


def test_pipeline_creates_ocr_xlsx_and_word_without_cross_route_fallback(tmp_path):
    config = Path("demo/projects/tongfu.yaml")
    template = Path("资产评估工作流/评估报告版式-沟通标注版.docx")
    pdf = Path("资产评估工作流/通富2025.6.30合并及母公司审计报告.pdf")
    template_hash = hashlib.sha256(template.read_bytes()).hexdigest()

    result = run_pipeline(
        project_config=config,
        pdf_path=pdf,
        output_dir=tmp_path,
        ocr_adapter=FixtureOcrAdapter(),
        llm_adapter=FixtureLlmAdapter(),
        qichacha_adapter=FixtureQichachaAdapter(),
        node_inputs={
            "selected_valuation_method": "收益法、资产基础法",
            "valuation_purpose_inputs": "用于股权收购决策。",
            "company_profile_section": "不应采用的节点越权内容",
        },
        ocr_field_resolver=fixture_ocr_fields,
        template_page_reader=FixtureTemplatePageReader(),
    )

    assert result.ocr_workbook_path.exists()
    assert result.report_path.exists()
    assert not (tmp_path / "字段审计清单.xlsx").exists()
    assert hashlib.sha256(template.read_bytes()).hexdigest() == template_hash
    fields = json.loads((tmp_path / "normalized_fields.json").read_text(encoding="utf-8"))
    assert fields["company_profile_section"] == "基于 OCR 证据生成的公司概述。"
    assert fields["main_products"] == "热处理服务。"
    assert fields["ownership_history"] == "企查查返回的历史股权沿革。"
    assert "所处行业及行业介绍" in fields["industry_overview"]
    assert fields["tax_rates"].startswith("增值税税率13%")
    assert fields["selected_valuation_method"] == "收益法、资产基础法"
    assert fields["commissioning_party_profile"] == ""
    required_monetary = json.loads(config.read_text(encoding="utf-8"))["required_monetary_fields"]
    assert required_monetary == [
        "registered_capital",
        "book_net_assets",
        "income_approach_value",
        "asset_approach_value",
        "historical_balance_sheet_table",
        "historical_income_statement_table",
        "major_long_term_assets",
        "asset_approach_result_section",
    ]
    assert all(fields.get(field) not in (None, "", [], {}) for field in required_monetary)
    assert "节点输入 返回越权字段，已丢弃：company_profile_section" in result.issues
    assert not any("commissioning_party_name" in issue for issue in result.issues)

    report = Document(result.report_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in report.paragraphs)
    assert "单体层面各类资产负债的金额为：货币资金" not in paragraph_text
    assert "单体层面各类资产负债的金额为：" in paragraph_text
    balance_text = "\n".join(cell.text for row in report.tables[6].rows for cell in row.cells)
    assert "148,537,259.26" in balance_text
    income_widths = [column.width for column in report.tables[5].columns]
    assert income_widths[0] < sum(income_widths[1:]) / 3 * 1.5
    ownership_text = "\n".join(
        cell.text for table_index in (2, 3) for row in report.tables[table_index].rows for cell in row.cells
    )
    assert "富士和机械工业（昆山）有限公司" not in ownership_text
    assert "上海上大热处理有限公司" not in ownership_text

    with zipfile.ZipFile(result.report_path) as archive:
        xml = "".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if re.fullmatch(r"word/(document|header\d+|footer\d+)\.xml", name)
        )
    assert 'w:val="yellow"' in xml
    assert "待人工补充" not in xml
    assert "不应采用" not in xml
    assert not (tmp_path / "生成问题清单.xlsx").exists()
    assert not (tmp_path / "生成问题清单.json").exists()

    manifest = json.loads(result.manifest_path.read_text(encoding="utf-8"))
    assert manifest["yellow_route_version"] == "yellow_routes.v1"
    assert manifest["prompt_version"] == "yellow_narratives.test"


def test_pipeline_does_not_run_llm_reviews_and_exports_four_node_trace(tmp_path):
    result = run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=Path("资产评估工作流/通富2025.6.30合并及母公司审计报告.pdf"),
        output_dir=tmp_path,
        ocr_adapter=FixtureOcrAdapter(),
        llm_adapter=FixtureLlmAdapter(),
        qichacha_adapter=FixtureQichachaAdapter(),
        ocr_field_resolver=fixture_ocr_fields,
        template_page_reader=FixtureTemplatePageReader(),
    )

    assert not (tmp_path / "格式审核.json").exists()
    assert not (tmp_path / "数据校验.json").exists()
    assert not (tmp_path / "语义审核.json").exists()
    assert not (tmp_path / "审核汇总.json").exists()
    manifest = json.loads((tmp_path / "run_manifest.json").read_text(encoding="utf-8"))
    assert "reviews" not in manifest
    trace_path = tmp_path / "workflow_trace.json"
    assert trace_path.exists()
    trace = json.loads(trace_path.read_text(encoding="utf-8"))
    assert trace["contract_version"] == "workflow_contract.v2"
    assert [node["node_name"] for node in trace["nodes"]] == [
        "start_input",
        "ocr_llm_candidates",
        "fill_word",
        "output",
    ]
    assert str(trace_path) in manifest["outputs"]


def test_pipeline_generates_review_report_when_monetary_fields_are_missing(
    tmp_path,
    monkeypatch,
):
    real_run_project = pipeline_module.run_project

    def run_project_with_missing_fields(*args, **kwargs):
        result = real_run_project(*args, **kwargs)
        fields_path = result.report_path.parent / "normalized_fields.json"
        fields = json.loads(fields_path.read_text(encoding="utf-8"))
        fields["book_net_assets"] = ""
        fields["historical_balance_sheet_table"] = ""
        fields_path.write_text(
            json.dumps(fields, ensure_ascii=False),
            encoding="utf-8",
        )
        return result

    monkeypatch.setattr(
        pipeline_module,
        "run_project",
        run_project_with_missing_fields,
    )

    result = run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=Path("资产评估工作流/通富2025.6.30合并及母公司审计报告.pdf"),
        output_dir=tmp_path,
        ocr_adapter=FixtureOcrAdapter(),
        llm_adapter=FixtureLlmAdapter(),
        qichacha_adapter=FixtureQichachaAdapter(),
        ocr_field_resolver=fixture_ocr_fields,
        template_page_reader=FixtureTemplatePageReader(),
    )

    assert result.report_path.exists()
    assert not (tmp_path / "资产评估报告_最终候选.docx").exists()
    assert "高优先级：金额及财务结果字段未匹配到，Word已保留黄色占位符：book_net_assets" in result.issues
    fields = json.loads(
        (tmp_path / "normalized_fields.json").read_text(encoding="utf-8")
    )
    assert fields["book_net_assets"] == ""
    report = Document(result.report_path)
    assert all(
        cell.text == "XXX"
        for row in report.tables[4].rows[1:]
        for cell in row.cells[1:]
    )

    manifest = json.loads(result.manifest_path.read_text(encoding="utf-8"))
    assert manifest["financial_validation"]["valid"] is False
    assert set(manifest["financial_validation"]["missing_fields"]) == {
        "book_net_assets",
        "historical_balance_sheet_table",
    }
    trace = json.loads(
        (tmp_path / "workflow_trace.json").read_text(encoding="utf-8")
    )
    fill_word_node = next(
        node for node in trace["nodes"] if node["node_name"] == "fill_word"
    )
    assert fill_word_node["status"] == "completed_with_issues"


def test_pipeline_without_pdf_exports_report_only(tmp_path):
    result = run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=None,
        output_dir=tmp_path,
        ocr_adapter=None,
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": None,
            "income_workbook": None,
            "reporting_workbook": None,
        },
        manual_inputs_override={"target_company_name": "示例有限公司"},
    )

    assert result.report_path.exists()
    assert result.ocr_workbook_path is None
    assert not (tmp_path / "生成问题清单.xlsx").exists()
    assert not (tmp_path / "生成问题清单.json").exists()
    manifest = json.loads(result.manifest_path.read_text(encoding="utf-8"))
    assert manifest["financial_validation"]["valid"] is False
    assert manifest["generation_validation"]["valid"] is False


def test_pipeline_preserves_semantic_excel_file_and_cell_evidence(tmp_path):
    reporting = tmp_path / "任意资产表.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "汇总表"
    sheet.append(["金额单位：人民币万元"])
    sheet.append(["项目", "账面价值", "评估价值"])
    sheet.append(["流动资产", 100, 110])
    sheet.append(["负债合计", 40, 42])
    sheet.append(["净资产", 60, 68])
    workbook.save(reporting)

    output = tmp_path / "run"
    result = run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=None,
        output_dir=output,
        ocr_adapter=None,
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": None,
            "income_workbook": None,
            "reporting_workbook": reporting,
        },
        manual_inputs_override={"target_company_name": "示例有限公司"},
    )

    trace = json.loads(
        (output / "workflow_trace.json").read_text(encoding="utf-8")
    )
    assert [node["node_name"] for node in trace["nodes"]] == [
        "start_input", "ocr_llm_candidates", "fill_word", "output"
    ]
    normalized_evidence = json.loads(
        (output / "normalized_evidence.json").read_text(encoding="utf-8")
    )
    assert normalized_evidence["asset_scope_summary_table"] == {
        "kind": "semantic_excel",
        "file": reporting.name,
        "locator": "汇总表!B3；汇总表!B4；汇总表!B5",
    }
    manifest = json.loads(
        (output / "run_manifest.json").read_text(encoding="utf-8")
    )
    assert str(output / "normalized_evidence.json") in manifest["outputs"]


def test_pipeline_keeps_unfinished_appraisal_reason_in_issue_list(
    tmp_path,
    monkeypatch,
):
    reporting = tmp_path / "尚未完成评估.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "汇总表"
    sheet.append(["金额单位：人民币万元"])
    sheet.append(["项目", "账面价值", "评估价值"])
    sheet.append(["资产总计", 100, 0])
    sheet.append(["负债合计", 40, 0])
    sheet.append(["净资产", 60, 0])
    workbook.save(reporting)

    def page_mapping(records, page_texts, paragraph_texts):
        location_ids = {
            str(item.get("location_id", ""))
            for item in records
            if isinstance(item, dict)
        }
        return (
            {"DOCUMENT-P0558-H01": 16}
            if "DOCUMENT-P0558-H01" in location_ids
            else {}
        )

    monkeypatch.setattr(
        pipeline_module,
        "map_location_pages",
        page_mapping,
    )
    output = tmp_path / "run"
    result = run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=None,
        output_dir=output,
        ocr_adapter=None,
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": None,
            "income_workbook": None,
            "reporting_workbook": reporting,
        },
        manual_inputs_override={"target_company_name": "示例有限公司"},
        template_page_reader=FixtureTemplatePageReader(),
    )

    assert any("尚未完成评估" in issue for issue in result.issues)


def test_pipeline_without_pdf_passes_material_fields_to_llm(tmp_path):
    output = tmp_path / "run"
    run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=None,
        output_dir=output,
        ocr_adapter=None,
        llm_adapter=NoPdfEvidenceLlmAdapter(),
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": None,
            "income_workbook": None,
            "reporting_workbook": None,
        },
        manual_inputs_override={
            "target_company_name": "示例有限公司",
            "narrative_modules": ["industry_overview"],
        },
    )

    fields = json.loads((output / "normalized_fields.json").read_text(encoding="utf-8"))
    assert fields["company_profile_section"] == "根据已上传结构化材料生成。"
    assert fields["industry_overview"] == ""


def test_pipeline_passes_uploaded_reference_word_to_llm(tmp_path):
    reference = tmp_path / "任意名称的参考报告.docx"
    document = Document()
    document.add_paragraph("4.2、主要产品")
    document.add_paragraph("该公司主营产品为工业滤波器，并面向工业自动化客户销售。")
    document.save(reference)
    output = tmp_path / "run"

    run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=None,
        output_dir=output,
        ocr_adapter=None,
        llm_adapter=ReferenceDocumentEvidenceLlmAdapter(),
        source_overrides={
            "audit_pdf": None,
            "reference_report": reference,
            "audited_financials": None,
            "income_workbook": None,
            "reporting_workbook": None,
        },
        manual_inputs_override={
            "target_company_name": "示例有限公司",
            "narrative_modules": ["main_products"],
        },
    )

    fields = json.loads((output / "normalized_fields.json").read_text(encoding="utf-8"))
    assert fields["main_products"] == "主营产品为工业滤波器。"


def test_qichacha_profile_supplies_missing_registered_capital(tmp_path):
    output = tmp_path / "run"
    run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=None,
        output_dir=output,
        ocr_adapter=None,
        qichacha_adapter=CapitalQichachaAdapter(),
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": None,
            "income_workbook": None,
            "reporting_workbook": None,
        },
        manual_inputs_override={"target_company_name": "示例有限公司"},
    )

    fields = json.loads((output / "normalized_fields.json").read_text(encoding="utf-8"))
    assert fields["registered_capital"] == "1,250万元"


def test_pipeline_keeps_semantically_matched_scope_table(tmp_path):
    reporting = tmp_path / "changed-layout.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "汇总表"
    sheet.append(["金额单位：人民币万元"])
    sheet.append(["项目", "序号", "账面价值", "评估价值"])
    sheet.append(["流动资产", 1, 100, 101])
    sheet.append(["非流动资产", 2, 200, 210])
    sheet.append(["固定资产", 21, 120, 130])
    sheet.append(["无形资产", 22, 30, 35])
    sheet.append(["长期待摊费用", 23, 10, 10])
    sheet.append(["资产总计", 3, 300, 311])
    sheet.append(["流动负债", 4, 50, 50])
    sheet.append(["非流动负债", 5, 20, 20])
    sheet.append(["负债总计", 6, 70, 70])
    sheet.append(["净资产", 7, 230, 241])
    detail = workbook.create_sheet("固定资产汇总表")
    detail.append(["金额单位：人民币元"])
    detail.append(["项目", "账面价值", "评估价值"])
    detail.append(["电子设备", 320_000, 350_000])
    workbook.save(reporting)

    run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=None,
        output_dir=tmp_path / "run",
        ocr_adapter=None,
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": reporting,
            "income_workbook": None,
            "reporting_workbook": reporting,
        },
        manual_inputs_override={"target_company_name": "示例有限公司"},
    )

    fields = json.loads(
        (tmp_path / "run/normalized_fields.json").read_text(encoding="utf-8")
    )
    rows = fields["asset_scope_summary_table"]["rows"]
    assert ["流动资产账面金额：", "1,000,000.00"] in rows
    assert ["所有者权益账面金额：", "2,300,000.00"] in rows
    long_term_text = "\n".join(
        cell.text
        for row in Document(tmp_path / "run/资产评估报告_待复核.docx").tables[7].rows
        for cell in row.cells
    )
    assert "320,000.00" in long_term_text


def test_pipeline_rejects_invalid_workflow_before_ocr(tmp_path):
    invalid_workflow = tmp_path / "invalid-workflow.json"
    invalid_workflow.write_text(
        json.dumps(
            {
                "version": "test",
                "contract_version": "workflow_contract.v1",
                "nodes": [
                    {
                        "name": "ocr_pdf",
                        "input_model": "MissingInput",
                        "output_model": "OcrPdfOutput",
                        "depends_on": [],
                        "human_checkpoint": None,
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    class FailIfCalledOcr:
        def extract(self, pdf_path):
            raise AssertionError("invalid workflow must stop before OCR")

    with pytest.raises(ValueError, match="工作流契约校验失败"):
        run_pipeline(
            project_config=Path("demo/projects/tongfu.yaml"),
            pdf_path=Path("资产评估工作流/通富2025.6.30合并及母公司审计报告.pdf"),
            output_dir=tmp_path / "run",
            ocr_adapter=FailIfCalledOcr(),
            workflow_path=invalid_workflow,
        )


def test_pipeline_keeps_review_report_when_cloud_ocr_fails(tmp_path):
    class FailingCloudOcr:
        def extract(self, _pdf_path):
            return [], ["阿里云 OCR 提交失败：NoPermission"]

    result = run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=Path(
            "资产评估工作流/通富2025.6.30合并及母公司审计报告.pdf"
        ),
        output_dir=tmp_path,
        ocr_adapter=FailingCloudOcr(),
        template_page_reader=FixtureTemplatePageReader(),
    )

    assert result.report_path.exists()
    assert "阿里云 OCR 提交失败：NoPermission" in result.issues
    assert not (tmp_path / "生成问题清单.xlsx").exists()


def test_pipeline_only_fills_selected_narrative_modules(tmp_path):
    result = run_pipeline(
        project_config=Path("demo/projects/tongfu.yaml"),
        pdf_path=Path("资产评估工作流/通富2025.6.30合并及母公司审计报告.pdf"),
        output_dir=tmp_path,
        ocr_adapter=FixtureOcrAdapter(),
        llm_adapter=FixtureLlmAdapter(),
        qichacha_adapter=FixtureQichachaAdapter(),
        ocr_field_resolver=fixture_ocr_fields,
        template_page_reader=FixtureTemplatePageReader(),
        manual_inputs_override={"narrative_modules": ["main_products"]},
    )

    fields = json.loads((tmp_path / "normalized_fields.json").read_text(encoding="utf-8"))
    assert fields["main_products"] == "热处理服务。"
    assert fields["industry_overview"] == ""
    assert fields["customers_suppliers"] == ""
