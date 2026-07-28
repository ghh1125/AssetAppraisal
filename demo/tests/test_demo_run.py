import hashlib
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook

from demo.run import run_project


class FakeLlm:
    def generate(self, evidence):
        assert evidence["target_company_name"]
        return {"company_profile_section": "由注入服务生成的公司简介"}, []


def test_offline_real_template_run_creates_new_report_and_audit(tmp_path: Path):
    config = Path("demo/projects/tongfu.yaml")
    template = Path("资产评估工作流/评估报告版式-沟通标注版.docx")
    before = hashlib.sha256(template.read_bytes()).hexdigest()
    result = run_project(config, output_dir=tmp_path, offline=True)
    assert hashlib.sha256(template.read_bytes()).hexdigest() == before
    assert result.report_path.exists() and result.report_path.resolve() != template.resolve()
    assert result.audit_path.exists()
    assert (tmp_path / "run_manifest.json").exists()
    assert (tmp_path / "issues.json").exists()
    with zipfile.ZipFile(result.report_path) as zf:
        xml = "".join(zf.read(name).decode("utf-8") for name in zf.namelist() if re.fullmatch(r"word/(document|header\d+|footer\d+)\.xml", name))
    assert "XXX" in xml
    assert 'w:highlight w:val="yellow"' in xml
    assert "企查查API获取" not in xml


def test_offline_run_leaves_missing_fields_blank_and_fills_material_assets(tmp_path: Path):
    result = run_project(Path("demo/projects/tongfu.yaml"), output_dir=tmp_path, offline=True)
    text = "\n".join(paragraph.text for paragraph in Document(result.report_path).paragraphs)
    fields = json.loads((tmp_path / "normalized_fields.json").read_text(encoding="utf-8"))
    assert "【待人工补充：" not in text
    assert "前海联合基金" not in text
    assert fields["valuation_scope"] == ""
    assert "三、评估对象和评估范围" in text
    assert "长期股权投资账面价值8,400,000.00元" not in text
    assert "固定资产账面净值4,993,561.04元" not in text
    assert "长期待摊费用账面价值1,788,311.49元" not in text
    assert "现有材料未见商标申报记录" in fields["trademark_summary"]
    assert "持有84项专利，其中在用专利63项" in text
    assert not fields["unrecorded_intangibles"].endswith("明细如下：")

    # The paragraph is only the table lead-in; all balance-sheet amounts
    # belong in the following configured asset/liability table.
    assert "单体层面各类资产负债的金额为：" in text
    assert "单体层面各类资产负债的金额为：货币资金" not in text
    balance_table = Document(result.report_path).tables[6]
    balance_text = "\n".join(cell.text for row in balance_table.rows for cell in row.cells)
    assert "流动资产账面金额" in balance_text
    assert "148,537,259.26" in balance_text
    long_asset_table = Document(result.report_path).tables[7]
    long_asset_text = "\n".join(cell.text for row in long_asset_table.rows for cell in row.cells)
    assert "419,017.57" in long_asset_text
    assert "1,788,311.49" in long_asset_text
    software_table = Document(result.report_path).tables[9]
    assert software_table.rows[0].cells[1].text == "软件名称"


def test_offline_run_fills_every_configured_financial_material_field(tmp_path: Path):
    config_path = Path("demo/projects/tongfu.yaml")
    result = run_project(config_path, output_dir=tmp_path, offline=True)
    config = json.loads(config_path.read_text(encoding="utf-8"))
    fields = json.loads((tmp_path / "normalized_fields.json").read_text(encoding="utf-8"))

    for key in config["required_financial_fields"]:
        assert fields[key] not in (None, "", []), key
        assert not any(issue.startswith(f"{key}：") for issue in result.issues)

    assert fields["registered_capital"] == "1,000.00万元"
    assert [fields[f"balance_history_year_{index}"] for index in range(1, 4)] == ["2023", "2024", "2025"]
    assert [fields[f"income_history_year_{index}"] for index in range(1, 4)] == ["2023", "2024", "2025"]
    assert fields["audit_report_name"] == "通富2025.6.30合并及母公司审计报告"
    assert "增值税税率为13%" in fields["tax_rates"]
    assert "企业所得税税率15%" in fields["tax_rates"]
    assert "热处理" in fields["main_products"]
    assert "6,365.04万元" in fields["asset_approach_result_section"]
    assert "1,766.88万元" in fields["asset_approach_result_section"]
    assert "38.43%" in fields["asset_approach_result_section"]

    report_text = "\n".join(paragraph.text for paragraph in Document(result.report_path).paragraphs)
    assert "、共同出资设立" not in report_text
    assert "有限公司、有限公司共同出资" not in report_text

    audit = load_workbook(result.audit_path, read_only=True, data_only=True)
    rows = list(audit["填充结果"].iter_rows(min_row=2, values_only=True))
    year_row = next(row for row in rows if row[4] == "balance_history_year_1")
    assert year_row[7] == "ocr_xlsx"
    assert "通富审核后财报-单体1月5日.xlsx" in year_row[8]
    assert "通富2025.6.30合并及母公司审计报告.pdf" in year_row[8]


def test_non_offline_run_uses_injected_provider_without_domain_dependency(tmp_path: Path):
    result = run_project(
        Path("demo/projects/tongfu.yaml"),
        output_dir=tmp_path,
        offline=False,
        report_date="2026-07-22",
        llm_adapter=FakeLlm(),
    )
    fields = json.loads((tmp_path / "normalized_fields.json").read_text(encoding="utf-8"))
    assert fields["company_profile_text"] == "由注入服务生成的公司简介"
    assert not any(issue.startswith("company_profile_text：") for issue in result.issues)


def test_missing_required_financial_field_does_not_block_review_report(tmp_path: Path):
    source_config_path = Path("demo/projects/tongfu.yaml").resolve()
    source_config = json.loads(source_config_path.read_text(encoding="utf-8"))
    source_base = source_config_path.parent
    for key in ("template", "mapping", "manual_inputs"):
        source_config[key] = str((source_base / source_config[key]).resolve())
    source_config["sources"] = {
        key: str((source_base / value).resolve())
        for key, value in source_config["sources"].items()
    }
    source_config["required_financial_fields"].append("synthetic_missing_amount")
    config_path = tmp_path / "missing-financial.json"
    config_path.write_text(
        json.dumps(source_config, ensure_ascii=False),
        encoding="utf-8",
    )

    output_dir = tmp_path / "run"
    result = run_project(config_path, output_dir=output_dir, offline=True)

    fields = json.loads(
        (output_dir / "normalized_fields.json").read_text(encoding="utf-8")
    )
    manifest = json.loads(result.manifest_path.read_text(encoding="utf-8"))
    assert result.report_path.exists()
    assert fields["synthetic_missing_amount"] == ""
    assert (
        "高优先级：财务材料字段未匹配到，Word已保留黄色占位符：synthetic_missing_amount"
        in result.issues
    )
    assert manifest["financial_validation"] == {
        "valid": False,
        "missing_fields": ["synthetic_missing_amount"],
        "conflicts": [],
    }


def test_run_project_generates_with_only_one_manual_field(tmp_path: Path):
    result = run_project(
        Path("demo/projects/tongfu.yaml"),
        output_dir=tmp_path,
        offline=True,
        manual_inputs_override={"target_company_name": "示例有限公司"},
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": None,
            "income_workbook": None,
            "reporting_workbook": None,
        },
    )

    assert result.report_path.exists()
    assert (tmp_path / "生成问题清单.xlsx").exists()
    assert (tmp_path / "生成问题清单.json").exists()
    document = Document(result.report_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "XXX" in text


def test_run_project_uses_semantic_excel_fallback_for_changed_layouts(tmp_path: Path):
    reporting = tmp_path / "任意名称-资产表.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "资产评估结果分类汇总表（元）"
    sheet.append(["金额单位：人民币元"])
    sheet.append(["序号", "科目名称", "账面价值", "评估价值"])
    sheet.append([1, "七、所有者权益（净资产）", 86_979_689.29, 93_972_005.69])
    workbook.save(reporting)

    income = tmp_path / "任意名称-收益表.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "净现金流计算表"
    sheet.append(["金额单位：元"])
    sheet.append(["股东全部权益价值", None, 68_500_000])
    workbook.save(income)

    output = tmp_path / "run"
    run_project(
        Path("demo/projects/tongfu.yaml"),
        output_dir=output,
        offline=True,
        manual_inputs_override={
            "target_company_name": "示例公司",
            "valuation_subject_type": "股东全部权益价值",
            "selected_valuation_method": "收益法、资产基础法",
            "final_valuation_method": "收益法",
            "transaction_type": "收购",
        },
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": reporting,
            "income_workbook": income,
            "reporting_workbook": reporting,
        },
    )

    fields = json.loads((output / "normalized_fields.json").read_text(encoding="utf-8"))
    evidence_path = output / "normalized_evidence.json"
    assert evidence_path.exists()
    evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
    assert fields["book_net_assets"] == 8697.968929
    assert fields["asset_approach_value"] == 9397.200569
    assert fields["income_approach_value"] == 6850
    assert fields["final_appraisal_value"] == 6850
    assert evidence["asset_approach_value"] == {
        "kind": "semantic_excel",
        "file": reporting.name,
        "locator": "资产评估结果分类汇总表（元）!D3",
    }
    assert evidence["income_approach_value"] == {
        "kind": "semantic_excel",
        "file": income.name,
        "locator": "净现金流计算表!C2",
    }


def test_run_project_prefers_complete_dated_history_from_later_workbook(tmp_path: Path):
    reporting = tmp_path / "资产基础法.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "资产负债表"
    sheet.append(["金额单位：人民币元"])
    sheet.append(["项目", "期初数", "期末数"])
    sheet.append(["资产总计", 200, 300])
    sheet.append(["负债合计", 70, 90])
    sheet.append(["所有者权益合计", 130, 210])
    workbook.save(reporting)

    income = tmp_path / "收益法.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "历资表"
    sheet.append(["金额单位：人民币元"])
    sheet.append(["项目", "2022年度", "2023年度", "2024年度"])
    sheet.append(["资产总计", 100, 200, 300])
    sheet.append(["负债合计", 40, 70, 90])
    sheet.append(["所有者权益合计", 60, 130, 210])
    workbook.save(income)

    output = tmp_path / "run"
    run_project(
        Path("demo/projects/tongfu.yaml"),
        output_dir=output,
        offline=True,
        manual_inputs_override={"target_company_name": "示例公司"},
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": None,
            "reporting_workbook": reporting,
            "income_workbook": income,
        },
    )

    fields = json.loads((output / "normalized_fields.json").read_text(encoding="utf-8"))
    assert fields["historical_balance_sheet_table"]["rows"] == [
        ["项目\\报表日", "2022年度", "2023年度", "2024年度"],
        ["总资产", "100.00", "200.00", "300.00"],
        ["负债", "40.00", "70.00", "90.00"],
        ["所有者权益", "60.00", "130.00", "210.00"],
    ]


def test_semantic_scope_table_replaces_readable_legacy_coordinates(tmp_path: Path):
    reporting = tmp_path / "新格式资产表.xlsx"
    workbook = Workbook()
    legacy = workbook.active
    legacy.title = "06N_资产负债表"
    for coordinate in (
        "F28",
        "F38",
        "F42",
        "F43",
        "F46",
        "F47",
        "F48",
        "F49",
        "F76",
        "L32",
        "L45",
        "L46",
        "L75",
    ):
        legacy[coordinate] = 999
    semantic = workbook.create_sheet("汇总表")
    semantic.append(["金额单位：人民币万元"])
    semantic.append(["项目", "账面价值", "评估价值"])
    semantic.append(["流动资产", 100, 110])
    semantic.append(["负债合计", 40, 42])
    semantic.append(["净资产", 60, 68])
    workbook.save(reporting)

    output = tmp_path / "run"
    run_project(
        Path("demo/projects/tongfu.yaml"),
        output_dir=output,
        offline=True,
        manual_inputs_override={"target_company_name": "示例公司"},
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": reporting,
            "reporting_workbook": reporting,
            "income_workbook": None,
        },
    )

    fields = json.loads((output / "normalized_fields.json").read_text(encoding="utf-8"))
    evidence = json.loads(
        (output / "normalized_evidence.json").read_text(encoding="utf-8")
    )
    assert fields["asset_scope_summary_table"]["rows"][0] == [
        "流动资产账面金额：",
        "1,000,000.00",
    ]
    assert evidence["asset_scope_summary_table"]["kind"] == "semantic_excel"
    assert evidence["asset_scope_summary_table"]["file"] == reporting.name
    assert "汇总表!B3" in evidence["asset_scope_summary_table"]["locator"]


def test_unfinished_appraisal_evidence_reaches_result_section(tmp_path: Path):
    reporting = tmp_path / "尚未完成评估.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "汇总表"
    sheet.append(["金额单位：人民币万元"])
    sheet.append(["项目", "账面价值", "评估价值"])
    sheet.append(["资产总计", 100, 0])
    sheet.append(["负债合计", 40, 0])
    sheet.append(["净资产", 60, 0])
    workbook.save(reporting)

    output = tmp_path / "run"
    run_project(
        Path("demo/projects/tongfu.yaml"),
        output_dir=output,
        offline=True,
        manual_inputs_override={"target_company_name": "示例公司"},
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": None,
            "reporting_workbook": reporting,
            "income_workbook": None,
        },
    )

    evidence = json.loads(
        (output / "normalized_evidence.json").read_text(encoding="utf-8")
    )
    assert evidence["asset_approach_result_section"] == {
        "kind": "unfinished_appraisal",
        "file": reporting.name,
        "locator": "汇总表!C5",
    }


def test_run_project_uses_selected_market_result_as_final_value(tmp_path: Path):
    market = tmp_path / "任意名称.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "汇总表"
    sheet.append(["评估方法", "市场法"])
    sheet.append(["金额单位：人民币万元"])
    sheet.append(["项目", "序号", "账面价值", "评估价值"])
    sheet.append(["净资产", 1, 29_151.74, 101_000])
    workbook.save(market)

    output = tmp_path / "run"
    run_project(
        Path("demo/projects/tongfu.yaml"),
        output_dir=output,
        offline=True,
        manual_inputs_override={
            "target_company_name": "示例公司",
            "valuation_subject_type": "股东全部权益价值",
            "selected_valuation_method": "市场法",
            "final_valuation_method": "市场法",
            "transaction_type": "收购",
        },
        source_overrides={
            "audit_pdf": None,
            "reference_report": None,
            "audited_financials": None,
            "income_workbook": market,
            "reporting_workbook": None,
        },
    )

    fields = json.loads((output / "normalized_fields.json").read_text(encoding="utf-8"))
    assert fields["market_approach_value"] == 101_000
    assert fields["final_appraisal_value"] == 101_000
