from pathlib import Path

from docx import Document

from demo.adapters.document import read_narrative_evidence


def test_narrative_evidence_prioritizes_late_relevant_content_within_limit(
    tmp_path: Path,
):
    report = tmp_path / "任意名称.docx"
    document = Document()
    for index in range(80):
        document.add_paragraph(f"普通说明段落{index}，不包含业务事实。")
    document.add_paragraph("4.2、主要产品")
    document.add_paragraph("公司主要生产工业滤波器和电抗器。")
    document.save(report)

    evidence = read_narrative_evidence(report, max_blocks=10)

    assert len(evidence) <= 10
    assert any("主要生产工业滤波器" in item["text"] for item in evidence)
